"""Pruebas de los cuatro exportadores."""

import json

import fitz
from docx import Document

from app.exportacion.datos_informe import (
    Informe,
    montos_desde_resumen,
    separar_secciones,
)
from app.exportacion.exportar_docx import exportar_docx
from app.exportacion.exportar_pdf import exportar_pdf
from app.exportacion.exportar_texto import exportar_json, exportar_txt

RESUMEN = """## Número y carátula del expediente
EXP-4567/2023 — Contratación de obra pública (pág. 1)

## Organismo iniciador
Ministerio de Obras Públicas (pág. 1)

## Montos
- $ 1.500.000 presupuesto oficial (pág. 12)
- $ 1.480.000 oferta adjudicada (pág. 30)

## Conclusión general
El trámite se encuentra en etapa de ejecución (pág. 45).
"""


def _informe():
    return Informe(
        titulo="Análisis de expediente",
        expediente="EXP-4567/2023",
        resumen_general=RESUMEN,
        resumenes_parciales=["### Bloque 1 (páginas 1–15)\n\n- Inicio del trámite (pág. 1)"],
        cronologia=[
            {"fecha_iso": "2023-01-05", "fecha_texto": "05/01/2023",
             "contexto": "nota de inicio", "pagina": 1},
        ],
        documentos=[
            {"tipo": "Resolución", "referencia": "RESOLUCIÓN N° 123/2023", "pagina": 3},
        ],
        metadatos={"fecha_analisis": "2026-07-11", "total_paginas": 45, "modelo": "llama3.2"},
    )


def test_separar_secciones():
    secciones = separar_secciones(RESUMEN)
    titulos = [t for t, _ in secciones]
    assert titulos[0] == "Número y carátula del expediente"
    assert "Montos" in titulos
    assert "EXP-4567/2023" in secciones[0][1]


def test_separar_secciones_sin_encabezados():
    secciones = separar_secciones("texto plano sin markdown")
    assert secciones == [("", "texto plano sin markdown")]


def test_montos_desde_resumen():
    montos = montos_desde_resumen(RESUMEN)
    assert len(montos) == 2
    assert "1.500.000" in montos[0]


def test_exportar_docx(tmp_path):
    ruta = tmp_path / "informe.docx"
    exportar_docx(_informe(), ruta)
    doc = Document(str(ruta))
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "Análisis de expediente" in texto
    assert "Resumen general del expediente" in texto
    assert "Índice" in texto
    # encabezado y tablas
    assert "EXP-4567/2023" in doc.sections[0].header.paragraphs[0].text
    tablas = doc.tables
    assert len(tablas) == 3  # cronología, montos, documentación
    assert "RESOLUCIÓN N° 123/2023" in tablas[2].rows[1].cells[1].text


def test_exportar_pdf(tmp_path):
    ruta = tmp_path / "informe.pdf"
    exportar_pdf(_informe(), ruta)
    with fitz.open(str(ruta)) as doc:
        assert doc.page_count >= 2  # carátula + contenido
        texto = "".join(p.get_text() for p in doc)
    assert "Análisis de expediente" in texto
    assert "Cronología del expediente" in texto
    assert "RESOLUCIÓN N° 123/2023" in texto


def test_exportar_txt(tmp_path):
    ruta = tmp_path / "informe.txt"
    exportar_txt(_informe(), ruta)
    texto = ruta.read_text(encoding="utf-8")
    assert "RESUMEN GENERAL" in texto
    assert "CRONOLOGÍA" in texto
    assert "(pág. 3)" in texto


def test_exportar_json(tmp_path):
    ruta = tmp_path / "informe.json"
    exportar_json(_informe(), ruta)
    datos = json.loads(ruta.read_text(encoding="utf-8"))
    assert datos["expediente"] == "EXP-4567/2023"
    assert datos["cronologia"][0]["fecha_iso"] == "2023-01-05"
    secciones = {s["titulo"] for s in datos["resumen_general"]["secciones"]}
    assert "Montos" in secciones
