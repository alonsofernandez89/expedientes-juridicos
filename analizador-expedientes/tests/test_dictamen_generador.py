"""Pruebas del ensamblado final del dictamen (.docx)."""

from docx import Document

from app.dictamen.cotejo import cotejar_requisitos
from app.dictamen.generador import datos_expediente_desde_resumen, generar_dictamen
from app.dictamen.requisitos import Requisito


def _resultados():
    requisitos = [
        Requisito(0, "Art. 5", "deberá constar informe técnico", ["informe", "tecnico"], "Informe técnico"),
        Requisito(1, "Art. 8", "es obligatorio el dictamen previo", ["dictamen", "previo"], "Dictamen"),
    ]
    documentos = [{"tipo": "Informe técnico", "referencia": "Informe 1", "pagina": 4}]
    return cotejar_requisitos(requisitos, paginas=["nada relevante"], documentos=documentos)


def test_generar_dictamen_incluye_secciones_y_tabla(tmp_path):
    ruta = tmp_path / "dictamen.docx"
    redaccion = {
        "considerando": "Que corresponde expedirse (pág. 4).\nQue falta el dictamen previo.",
        "resuelve": "Art. 1°. Continuar el trámite.\nArt. 2°. Notificar.",
    }
    generar_dictamen(
        titulos_secciones=["VISTO", "CONSIDERANDO", "RESUELVE"],
        resultados=_resultados(),
        redaccion=redaccion,
        datos_expediente={"caratula": "EXP-1/2024", "organismo": "Ministerio X", "objeto": "compra de insumos"},
        ruta_salida=ruta,
    )
    assert ruta.exists()
    doc = Document(str(ruta))
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "DICTAMEN JURÍDICO" in texto
    assert "EXP-1/2024" in texto
    assert "Que corresponde expedirse" in texto
    assert "Art. 1°. Continuar el trámite." in texto
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[1].cells[0].text == "Art. 5"


def test_generar_dictamen_seccion_desconocida_queda_marcada(tmp_path):
    ruta = tmp_path / "dictamen.docx"
    generar_dictamen(
        titulos_secciones=["ANTECEDENTES"],
        resultados=_resultados(),
        redaccion={"considerando": "", "resuelve": ""},
        datos_expediente={},
        ruta_salida=ruta,
    )
    doc = Document(str(ruta))
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "completar manualmente" in texto


def test_datos_expediente_desde_resumen():
    resumen = """## Número y carátula del expediente
EXP-9/2023 — Compra directa

## Organismo iniciador
Dirección de Compras

## Objeto de las actuaciones
Adquisición de mobiliario
"""
    datos = datos_expediente_desde_resumen(resumen)
    assert datos["caratula"] == "EXP-9/2023 — Compra directa"
    assert datos["organismo"] == "Dirección de Compras"
    assert datos["objeto"] == "Adquisición de mobiliario"


def test_datos_expediente_desde_resumen_vacio():
    assert datos_expediente_desde_resumen("") == {}
