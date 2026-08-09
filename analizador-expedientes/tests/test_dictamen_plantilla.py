"""Pruebas de lectura de plantilla de dictamen y de normativa."""

from docx import Document

from app.dictamen.plantilla import (
    ESTRUCTURA_POR_DEFECTO,
    clasificar_seccion,
    leer_estructura_plantilla,
    leer_texto_normativa,
)


def _crear_plantilla(ruta, titulos_con_estilo=True):
    doc = Document()
    if titulos_con_estilo:
        doc.add_heading("VISTO", level=1)
        doc.add_paragraph("El expediente N° ___ caratulado ___.")
        doc.add_heading("CONSIDERANDO", level=1)
        doc.add_paragraph("Que corresponde expedirse sobre el trámite.")
        doc.add_heading("RESUELVE", level=1)
        doc.add_paragraph("Artículo 1°. Aprobar lo actuado.")
    else:
        # títulos como líneas cortas en mayúsculas, sin estilo Heading
        doc.add_paragraph("VISTO")
        doc.add_paragraph("El expediente caratulado como se indica.")
        doc.add_paragraph("CONSIDERANDO")
        doc.add_paragraph("Que corresponde expedirse.")
        doc.add_paragraph("RESUELVE")
        doc.add_paragraph("Aprobar lo actuado.")
    doc.save(str(ruta))


def test_leer_estructura_con_estilos_heading(tmp_path):
    ruta = tmp_path / "plantilla.docx"
    _crear_plantilla(ruta, titulos_con_estilo=True)
    titulos = leer_estructura_plantilla(ruta)
    assert titulos == ["VISTO", "CONSIDERANDO", "RESUELVE"]


def test_leer_estructura_por_lineas_mayusculas(tmp_path):
    ruta = tmp_path / "plantilla.docx"
    _crear_plantilla(ruta, titulos_con_estilo=False)
    titulos = leer_estructura_plantilla(ruta)
    assert titulos == ["VISTO", "CONSIDERANDO", "RESUELVE"]


def test_leer_estructura_plantilla_sin_titulos_usa_default(tmp_path):
    ruta = tmp_path / "plantilla.docx"
    doc = Document()
    doc.add_paragraph("Un párrafo común y silvestre sin ningún título especial.")
    doc.save(str(ruta))
    assert leer_estructura_plantilla(ruta) == ESTRUCTURA_POR_DEFECTO


def test_clasificar_seccion_reconoce_sinonimos():
    assert clasificar_seccion("VISTO") == "visto"
    assert clasificar_seccion("Vistos:") == "visto"
    assert clasificar_seccion("Y CONSIDERANDO") == "considerando"
    assert clasificar_seccion("SE RESUELVE") == "resuelve"
    assert clasificar_seccion("DICTAMINA") == "resuelve"
    assert clasificar_seccion("ANTECEDENTES") is None


def test_leer_texto_normativa_txt(tmp_path):
    ruta = tmp_path / "norma.txt"
    ruta.write_text("Artículo 1°. El interesado deberá presentar nota.", encoding="utf-8")
    assert "deberá presentar nota" in leer_texto_normativa(ruta)


def test_leer_texto_normativa_docx(tmp_path):
    ruta = tmp_path / "norma.docx"
    doc = Document()
    doc.add_paragraph("Artículo 2°. Se requiere presupuesto oficial.")
    doc.save(str(ruta))
    assert "presupuesto oficial" in leer_texto_normativa(ruta)


def test_leer_texto_normativa_pdf(tmp_path):
    import fitz

    ruta = tmp_path / "norma.pdf"
    doc = fitz.open()
    pagina = doc.new_page()
    pagina.insert_text((72, 72), "Articulo 3. Es obligatorio el dictamen previo.")
    doc.save(str(ruta))
    doc.close()
    assert "dictamen previo" in leer_texto_normativa(ruta)
