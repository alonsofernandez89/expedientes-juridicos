"""Exportación a Word (.docx) con formato profesional.

Incluye: carátula, índice automático (campo TOC), encabezado con el número
de expediente, pie con numeración de páginas, subtítulos con estilos, y
tablas para fechas relevantes, montos y documentación detectada.

El índice usa un campo TOC de Word: al abrir el archivo, Word pide
actualizar campos (o F9) y el índice se completa con las páginas reales.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.exportacion.datos_informe import (
    Informe,
    montos_desde_resumen,
    separar_secciones,
)

_COLOR_TITULO = RGBColor(0x1F, 0x3B, 0x5C)


def _campo(parrafo, instruccion: str) -> None:
    """Inserta un campo de Word (PAGE, TOC, etc.) en el párrafo."""
    run = parrafo.add_run()
    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruccion
    separador = OxmlElement("w:fldChar")
    separador.set(qn("w:fldCharType"), "separate")
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    for el in (inicio, instr, separador, fin):
        run._r.append(el)


def _tabla(doc, encabezados: list[str], filas: list[list[str]]) -> None:
    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = "Light Grid Accent 1"
    for i, texto in enumerate(encabezados):
        celda = tabla.rows[0].cells[i]
        celda.text = texto
        for p in celda.paragraphs:
            for r in p.runs:
                r.bold = True
    for fila in filas:
        celdas = tabla.add_row().cells
        for i, valor in enumerate(fila):
            celdas[i].text = str(valor)


def _texto_markdown_simple(doc, texto: str) -> None:
    """Vuelca texto Markdown básico (listas y párrafos) al documento."""
    for linea in texto.splitlines():
        limpia = linea.strip()
        if not limpia:
            continue
        if limpia.startswith(("-", "*", "•")):
            doc.add_paragraph(limpia.lstrip("-*• ").strip(), style="List Bullet")
        else:
            doc.add_paragraph(limpia.replace("**", ""))


def exportar_docx(informe: Informe, ruta: Path) -> None:
    doc = Document()

    # Estilos base
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for nivel, tamano in (("Heading 1", 16), ("Heading 2", 13)):
        estilo = doc.styles[nivel]
        estilo.font.color.rgb = _COLOR_TITULO
        estilo.font.size = Pt(tamano)

    seccion = doc.sections[0]
    seccion.top_margin = seccion.bottom_margin = Cm(2.5)
    seccion.left_margin = seccion.right_margin = Cm(2.5)

    # Encabezado: nombre del expediente
    p_enc = seccion.header.paragraphs[0]
    p_enc.text = f"Expediente: {informe.expediente}"
    p_enc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p_enc.runs:
        r.font.size = Pt(9)

    # Pie: leyenda + numeración de página (campo PAGE de Word)
    p_pie = seccion.footer.paragraphs[0]
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie.add_run("Análisis generado localmente — Página ").font.size = Pt(9)
    _campo(p_pie, "PAGE")
    p_pie.add_run(" de ").font.size = Pt(9)
    _campo(p_pie, "NUMPAGES")

    # --- Carátula ---
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = t.add_run(informe.titulo)
    rt.font.size = Pt(28)
    rt.bold = True
    rt.font.color.rgb = _COLOR_TITULO
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(informe.expediente)
    rs.font.size = Pt(18)
    for clave in ("fecha_analisis", "total_paginas", "modelo"):
        if clave in informe.metadatos:
            m = doc.add_paragraph()
            m.alignment = WD_ALIGN_PARAGRAPH.CENTER
            etiquetas = {
                "fecha_analisis": "Fecha de análisis",
                "total_paginas": "Páginas del expediente",
                "modelo": "Modelo local utilizado",
            }
            m.add_run(f"{etiquetas[clave]}: {informe.metadatos[clave]}").font.size = Pt(12)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- Índice ---
    doc.add_heading("Índice", level=1)
    aviso = doc.add_paragraph()
    r_aviso = aviso.add_run(
        "(En Word: clic derecho sobre el índice → «Actualizar campos» para "
        "completar los números de página.)"
    )
    r_aviso.font.size = Pt(9)
    r_aviso.italic = True
    _campo(doc.add_paragraph(), r'TOC \o "1-2" \h \z \u')
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- Resumen general por secciones ---
    doc.add_heading("Resumen general del expediente", level=1)
    for titulo, cuerpo in separar_secciones(informe.resumen_general):
        if titulo:
            doc.add_heading(titulo, level=2)
        if cuerpo:
            _texto_markdown_simple(doc, cuerpo)

    # --- Tabla de fechas relevantes (cronología) ---
    if informe.cronologia:
        doc.add_heading("Cronología del expediente", level=1)
        _tabla(
            doc,
            ["Fecha", "Hecho / contexto", "Página"],
            [
                [e.get("fecha_texto") or e.get("fecha_iso", ""),
                 e.get("contexto", ""), e.get("pagina", "")]
                for e in informe.cronologia
            ],
        )

    # --- Tabla de montos ---
    montos = montos_desde_resumen(informe.resumen_general)
    if montos:
        doc.add_heading("Montos detectados", level=1)
        _tabla(doc, ["Monto / concepto"], [[m] for m in montos])

    # --- Tabla de documentación detectada ---
    if informe.documentos:
        doc.add_heading("Documentación detectada", level=1)
        _tabla(
            doc,
            ["Tipo", "Referencia", "Página"],
            [[d["tipo"], d["referencia"], d["pagina"]] for d in informe.documentos],
        )

    # --- Resúmenes parciales (anexo) ---
    if informe.resumenes_parciales:
        doc.add_heading("Anexo: resúmenes por bloque", level=1)
        for parcial in informe.resumenes_parciales:
            for titulo, cuerpo in separar_secciones(parcial):
                if titulo:
                    doc.add_heading(titulo, level=2)
                if cuerpo:
                    _texto_markdown_simple(doc, cuerpo)

    ruta.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(ruta))
