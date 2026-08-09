"""Lectura de la plantilla de dictamen (.docx) y de la normativa.

De la plantilla sólo se extraen los TÍTULOS de sección, en el orden en que
aparecen (por ejemplo VISTO / CONSIDERANDO / RESUELVE, o los que use el
organismo). El dictamen generado reutiliza esa misma estructura.
"""

from pathlib import Path

LARGO_MAXIMO_TITULO = 60
PALABRAS_MAXIMO_TITULO = 8

ESTRUCTURA_POR_DEFECTO = ["VISTO", "CONSIDERANDO", "RESUELVE"]

# Sinónimos habituales de cada sección estándar de un dictamen administrativo
SINONIMOS_SECCION = {
    "visto": {"visto", "vistos"},
    "considerando": {"considerando", "considerandos", "y considerando"},
    "resuelve": {
        "resuelve", "se resuelve", "dictamina", "por ello", "por lo expuesto",
        "conclusion", "conclusión",
    },
}


def _es_titulo_corto(texto: str) -> bool:
    return (
        0 < len(texto) <= LARGO_MAXIMO_TITULO
        and len(texto.split()) <= PALABRAS_MAXIMO_TITULO
        and texto == texto.upper()
        and any(c.isalpha() for c in texto)
    )


def leer_estructura_plantilla(ruta: Path) -> list[str]:
    """Títulos de sección detectados en el .docx, en orden de aparición.

    Se consideran título tanto los párrafos con estilo "Heading" como las
    líneas cortas en mayúsculas (forma habitual de los dictámenes: VISTO,
    CONSIDERANDO, RESUELVE). Si no se detecta ninguno, se usa la estructura
    estándar de un dictamen administrativo.
    """
    from docx import Document

    doc = Document(str(ruta))
    titulos = []
    for parrafo in doc.paragraphs:
        texto = parrafo.text.strip()
        if not texto:
            continue
        es_encabezado = parrafo.style.name.lower().startswith("heading")
        if es_encabezado or _es_titulo_corto(texto):
            titulos.append(texto)
    return titulos or list(ESTRUCTURA_POR_DEFECTO)


def clasificar_seccion(titulo: str) -> str | None:
    """Mapea un título de sección a su rol estándar ('visto', 'considerando',
    'resuelve') o None si no se reconoce (se copia tal cual al dictamen)."""
    normalizado = titulo.strip().lower().rstrip(":.")
    for rol, sinonimos in SINONIMOS_SECCION.items():
        if normalizado in sinonimos:
            return rol
    return None


def leer_texto_normativa(ruta: Path) -> str:
    """Extrae el texto de la normativa desde PDF, DOCX o TXT."""
    sufijo = Path(ruta).suffix.lower()
    if sufijo == ".pdf":
        from app.extraccion.extractor import extraer_paginas

        return "\n\n".join(extraer_paginas(Path(ruta)))
    if sufijo in (".docx", ".doc"):
        from docx import Document

        doc = Document(str(ruta))
        return "\n".join(p.text for p in doc.paragraphs)
    return Path(ruta).read_text(encoding="utf-8", errors="ignore")
