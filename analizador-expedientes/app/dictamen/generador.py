"""Ensamblado del dictamen final (.docx) a partir de:

- la ESTRUCTURA de la plantilla cargada (títulos de sección, en orden);
- el COTEJO determinístico de requisitos (siempre incluido, en una tabla,
  para que el resultado sea auditable más allá de la prosa generada);
- la REDACCIÓN de Ollama para las secciones reconocidas (VISTO,
  CONSIDERANDO, RESUELVE); las secciones no reconocidas se copian con un
  aviso para completar a mano.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.dictamen.cotejo import CUMPLE, FALTA, ResultadoCotejo, resumen_cumplimiento
from app.dictamen.plantilla import clasificar_seccion

_COLOR_TITULO = RGBColor(0x1F, 0x3B, 0x5C)

_ETIQUETA_CORTA = {CUMPLE: "Cumple", "no_consta": "No consta", FALTA: "Falta"}


def _tabla_cumplimiento(doc: Document, resultados: list[ResultadoCotejo]) -> None:
    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = "Light Grid Accent 1"
    encabezados = tabla.rows[0].cells
    for i, texto in enumerate(["Artículo", "Requisito", "Estado", "Página"]):
        encabezados[i].text = texto
        for p in encabezados[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for r in resultados:
        celdas = tabla.add_row().cells
        celdas[0].text = r.requisito.articulo
        celdas[1].text = r.requisito.texto
        celdas[2].text = _ETIQUETA_CORTA.get(r.estado, r.estado)
        celdas[3].text = str(r.pagina) if r.pagina else "—"


def _parrafo_visto(doc: Document, datos_expediente: dict) -> None:
    partes = []
    if datos_expediente.get("caratula"):
        partes.append(f"El expediente {datos_expediente['caratula']}")
    else:
        partes.append("El expediente de referencia")
    if datos_expediente.get("organismo"):
        partes.append(f"iniciado por {datos_expediente['organismo']}")
    if datos_expediente.get("objeto"):
        partes.append(f"cuyo objeto es {datos_expediente['objeto']}")
    doc.add_paragraph(", ".join(partes) + ".")


def generar_dictamen(
    titulos_secciones: list[str],
    resultados: list[ResultadoCotejo],
    redaccion: dict,
    datos_expediente: dict,
    ruta_salida: Path,
) -> None:
    """Arma y guarda el dictamen final.

    `datos_expediente`: dict con claves opcionales caratula/organismo/objeto
    (se completan con lo ya extraído del resumen general, si existe).
    `redaccion`: {"considerando": str, "resuelve": str} (de fundamentacion.py).
    """
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("DICTAMEN JURÍDICO")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = _COLOR_TITULO
    doc.add_paragraph()

    conteo = resumen_cumplimiento(resultados)
    resumen = doc.add_paragraph()
    resumen.add_run(
        f"Requisitos evaluados: {conteo['total']}  ·  "
        f"Cumplen: {conteo[CUMPLE]}  ·  "
        f"No constan: {conteo['no_consta']}  ·  "
        f"Faltan: {conteo[FALTA]}"
    ).italic = True

    for titulo_seccion in titulos_secciones:
        doc.add_heading(titulo_seccion, level=1)
        rol = clasificar_seccion(titulo_seccion)
        if rol == "visto":
            _parrafo_visto(doc, datos_expediente)
        elif rol == "considerando":
            for parrafo in redaccion.get("considerando", "").split("\n"):
                if parrafo.strip():
                    doc.add_paragraph(parrafo.strip())
            doc.add_paragraph()
            doc.add_paragraph().add_run("Detalle del cotejo normativo:").bold = True
            _tabla_cumplimiento(doc, resultados)
        elif rol == "resuelve":
            for parrafo in redaccion.get("resuelve", "").split("\n"):
                if parrafo.strip():
                    doc.add_paragraph(parrafo.strip())
        else:
            aviso = doc.add_paragraph()
            aviso.add_run(
                "[Sección no generada automáticamente — completar manualmente]"
            ).italic = True

    Path(ruta_salida).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(ruta_salida))


def datos_expediente_desde_resumen(resumen_general: str) -> dict:
    """Extrae carátula/organismo/objeto del resumen general ya generado
    (mismo formato Markdown que usa la exportación de informes)."""
    from app.exportacion.datos_informe import separar_secciones

    mapa = {
        "número y carátula del expediente": "caratula",
        "organismo iniciador": "organismo",
        "objeto de las actuaciones": "objeto",
    }
    datos: dict = {}
    for titulo, cuerpo in separar_secciones(resumen_general or ""):
        clave = mapa.get(titulo.strip().lower())
        if clave and cuerpo.strip():
            datos[clave] = cuerpo.strip().splitlines()[0]
    return datos
